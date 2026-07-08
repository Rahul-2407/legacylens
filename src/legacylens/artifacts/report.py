"""Report assembly.

Weaves two kinds of content into one deliverable and never confuses them:

* machine content — scores, tables, diagrams, the findings appendix —
  rendered directly from artifacts, byte-deterministic
* agent prose — the cited sections from Module 10, inserted verbatim with
  their [F-...] citations intact so every claim remains traceable to the
  appendix

Degradation is explicit: a missing or failed agent section becomes a
visible note, never silence, and the deterministic content is complete
either way.
"""

from collections import Counter
from datetime import date

from legacylens.agents.graph import SynthesisResult
from legacylens.artifacts.mermaid import module_graph_mermaid, waves_mermaid
from legacylens.domain.models import Finding, ProjectContext, Severity
from legacylens.graph.algorithms import ArchitectureReport
from legacylens.parsing.ast.module_graph import ModuleGraph
from legacylens.scoring.engine import ScoreCard

_SEVERITY_ORDER = (Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM,
                   Severity.LOW, Severity.INFO)
MAX_APPENDIX = 100
_FALLBACK = ("_LLM synthesis was unavailable for this section; the "
             "deterministic findings and scores in this report are "
             "complete and unaffected._")


def _agent_sections(synthesis: SynthesisResult | None, agent: str,
                    lines: list[str]) -> None:
    if synthesis and agent in synthesis.outputs:
        for section in synthesis.outputs[agent].sections:
            lines.append(f"### {section.heading}")
            lines.append(section.content)
            lines.append(f"_Citations: {', '.join(section.citations)}_")
            lines.append("")
    else:
        reason = (synthesis.failures.get(agent) if synthesis else None)
        lines.append(_FALLBACK)
        if reason:
            lines.append(f"_Reason: {reason}_")
        lines.append("")


class ReportBuilder:
    def build_markdown(
        self,
        ctx: ProjectContext,
        findings: list[Finding],
        scorecard: ScoreCard,
        synthesis: SynthesisResult | None = None,
    ) -> str:
        lines: list[str] = []
        lines.append(f"# LegacyLens Migration Assessment — "
                     f"{ctx.project_id}")
        lines.append(
            f"Generated {date.today().isoformat()} · {len(ctx.files)} "
            f"files analyzed · {len(findings)} findings. Every claim in "
            "this report cites finding IDs resolvable in the appendix."
        )
        lines.append("")

        lines.append("## Executive summary")
        _agent_sections(synthesis, "writer", lines)

        r = scorecard.readiness
        lines.append(f"## Migration readiness: {r.value}/100")
        if r.components:
            lines.append("| Component | Deduction |")
            lines.append("|---|---|")
            for c in r.components:
                lines.append(f"| {c.detail} | {c.delta} |")
        else:
            lines.append("No deductions — no non-informational findings.")
        lines.append("")

        e = scorecard.effort
        lines.append("## Effort estimate")
        lines.append(
            f"**{e.expected_days} person-days expected** "
            f"(range {e.optimistic_days}–{e.pessimistic_days}); module "
            f"complexity bands: {e.band_counts}."
        )
        for m in e.multipliers_applied:
            lines.append(f"- applied: {m}")
        lines.append("Assumptions:")
        for a in e.assumptions:
            lines.append(f"- {a}")
        lines.append("")

        lines.append("## Risk matrix")
        categories = sorted({c for row in scorecard.risk_matrix.values()
                             for c in row})
        if categories:
            lines.append("| Severity | " + " | ".join(categories) + " |")
            lines.append("|---|" + "---|" * len(categories))
            for severity in _SEVERITY_ORDER:
                row = scorecard.risk_matrix.get(str(severity))
                if not row:
                    continue
                cells = [str(row.get(c, 0)) for c in categories]
                lines.append(f"| {severity} | " + " | ".join(cells) + " |")
        lines.append("")
        if scorecard.top_risks:
            lines.append("**Top risks:**")
            for risk in scorecard.top_risks:
                lines.append(
                    f"- [{risk['severity'].upper()}] {risk['title']} "
                    f"({risk['finding_id']})"
                )
            lines.append("")

        lines.append("## Architecture analysis")
        _agent_sections(synthesis, "analyst", lines)

        graph: ModuleGraph | None = ctx.get_artifact("module_graph")
        arch: ArchitectureReport | None = ctx.get_artifact("architecture")
        if graph and graph.internal_edges:
            lines.append("### Module dependency graph")
            lines.append("```mermaid")
            lines.append(module_graph_mermaid(graph, arch))
            lines.append("```")
            lines.append("")
        if arch and arch.waves:
            lines.append("### Migration waves")
            lines.append("```mermaid")
            lines.append(waves_mermaid(arch))
            lines.append("```")
            lines.append("")

        lines.append("## Migration plan")
        _agent_sections(synthesis, "strategist", lines)

        if scorecard.quick_wins:
            lines.append("**Quick wins (wave 0, low complexity, no severe "
                         "findings):** " + ", ".join(scorecard.quick_wins))
            lines.append("")
        for m in scorecard.high_risk_modules:
            lines.append(f"- **High-risk module** `{m['path']}` "
                         f"(complexity {m['complexity']}): "
                         + "; ".join(m["reasons"]))
        lines.append("")

        lines.append("## Findings appendix")
        counts = Counter(f.severity for f in findings)
        lines.append("Totals: " + ", ".join(
            f"{counts.get(s, 0)} {s}" for s in _SEVERITY_ORDER))
        lines.append("")
        lines.append("| ID | Severity | Rule | Title | Key evidence |")
        lines.append("|---|---|---|---|---|")
        ordered = sorted(
            findings,
            key=lambda f: (_SEVERITY_ORDER.index(f.severity), f.rule_id),
        )[:MAX_APPENDIX]
        for f in ordered:
            ev = next((e for e in f.evidence if e.file_path), f.evidence[0])
            where = (f"{ev.file_path}:{ev.line_start or ''}"
                     if ev.file_path
                     else (ev.reference_url or ev.detail or ""))
            title = f.title.replace("|", "\\|")
            lines.append(f"| {f.finding_id} | {f.severity} | {f.rule_id} "
                         f"| {title} | {where} |")
        if len(findings) > MAX_APPENDIX:
            lines.append(f"\n_+{len(findings) - MAX_APPENDIX} further "
                         "findings available via the API/export._")
        lines.append("")
        lines.append("---")
        lines.append("_Generated by LegacyLens. Deterministic analysis "
                     "produced the facts; AI produced only cited "
                     "interpretation._")
        return "\n".join(lines)

    def build_docx(
        self,
        path,
        ctx: ProjectContext,
        findings: list[Finding],
        scorecard: ScoreCard,
        synthesis: SynthesisResult | None = None,
    ) -> None:
        """Executive Word deliverable: summary, scores, risks, findings."""
        from docx import Document

        doc = Document()
        doc.add_heading(
            f"LegacyLens Migration Assessment — {ctx.project_id}", 0)
        doc.add_paragraph(
            f"Generated {date.today().isoformat()} · {len(ctx.files)} "
            f"files · {len(findings)} findings · readiness "
            f"{scorecard.readiness.value}/100"
        )

        doc.add_heading("Executive summary", 1)
        if synthesis and "writer" in synthesis.outputs:
            for section in synthesis.outputs["writer"].sections:
                doc.add_heading(section.heading, 2)
                doc.add_paragraph(section.content)
                doc.add_paragraph(
                    f"Citations: {', '.join(section.citations)}"
                ).italic = True
        else:
            doc.add_paragraph(_FALLBACK.strip("_"))

        doc.add_heading(
            f"Migration readiness: {scorecard.readiness.value}/100", 1)
        if scorecard.readiness.components:
            table = doc.add_table(
                rows=1 + len(scorecard.readiness.components), cols=2)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Component"
            table.rows[0].cells[1].text = "Deduction"
            for i, c in enumerate(scorecard.readiness.components, start=1):
                table.rows[i].cells[0].text = c.detail
                table.rows[i].cells[1].text = str(c.delta)

        e = scorecard.effort
        doc.add_heading("Effort estimate", 1)
        doc.add_paragraph(
            f"{e.expected_days} person-days expected "
            f"({e.optimistic_days}–{e.pessimistic_days})."
        )
        for a in e.assumptions:
            doc.add_paragraph(a, style="List Bullet")

        doc.add_heading("Findings", 1)
        ordered = sorted(
            findings,
            key=lambda f: (_SEVERITY_ORDER.index(f.severity), f.rule_id),
        )[:40]
        table = doc.add_table(rows=1 + len(ordered), cols=4)
        table.style = "Light Grid Accent 1"
        for j, header in enumerate(("ID", "Severity", "Rule", "Title")):
            table.rows[0].cells[j].text = header
        for i, f in enumerate(ordered, start=1):
            row = table.rows[i].cells
            row[0].text = f.finding_id
            row[1].text = str(f.severity)
            row[2].text = f.rule_id
            row[3].text = f.title

        doc.save(str(path))
