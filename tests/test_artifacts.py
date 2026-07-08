"""Artifact tests: mermaid validity and caps, markdown assembly with
citations intact, honest degradation notes, and a reopenable docx."""

from pathlib import Path

import pytest
from docx import Document

from legacylens.agents.citation import SectionedOutput
from legacylens.agents.graph import SynthesisResult
from legacylens.artifacts.mermaid import module_graph_mermaid, waves_mermaid
from legacylens.artifacts.report import ReportBuilder
from legacylens.domain.models import (
    Evidence,
    FileRecord,
    Finding,
    FindingCategory,
    ProjectContext,
    Severity,
)
from legacylens.graph.algorithms import analyze_architecture
from legacylens.parsing.ast.module_graph import ImportEdge, ModuleGraph
from legacylens.scoring.engine import compute_scorecard


def make_graph(edges):
    files = sorted({n for e in edges for n in e})
    return ModuleGraph(files=files, edges=[
        ImportEdge(source=s, target=t, raw=t, line=1, internal=True)
        for s, t in edges
    ])


@pytest.fixture()
def project(tmp_path):
    graph = make_graph([
        ("app/main.py", "app/orders.py"),
        ("app/orders.py", "app/billing.py"),
        ("app/billing.py", "app/orders.py"),      # cycle
        ("app/orders.py", "app/util.py"),
    ])
    arch = analyze_architecture(graph)
    findings = [
        Finding(analyzer_id="x", rule_id="TECH-EOL-001",
                category=FindingCategory.TECHNOLOGY,
                severity=Severity.CRITICAL,
                title="Django 1.11 is past end-of-life",
                description="d",
                evidence=[Evidence(file_path="requirements.txt",
                                   line_start=1)]),
        Finding(analyzer_id="x", rule_id="ARCH-CYCLE-001",
                category=FindingCategory.ARCHITECTURE,
                severity=Severity.HIGH, title="Cyclic dependency",
                description="d",
                evidence=[Evidence(file_path="app/orders.py",
                                   line_start=1)]),
    ]
    ctx = ProjectContext(
        project_id="proj-demo", root=tmp_path,
        files=[FileRecord(path=p, size_bytes=100) for p in graph.files],
    )
    ctx.artifacts["module_graph"] = graph
    ctx.artifacts["architecture"] = arch
    scorecard = compute_scorecard(findings, ctx)
    return ctx, findings, scorecard, graph, arch


def synthesis_for(findings, include=("analyst", "strategist", "writer"),
                  failures=None) -> SynthesisResult:
    fid = findings[0].finding_id
    outputs = {
        agent: SectionedOutput(sections=[{
            "heading": f"{agent} heading",
            "content": f"{agent} grounded prose citing [{fid}].",
            "citations": [fid],
        }])
        for agent in include
    }
    return SynthesisResult(outputs=outputs, failures=failures or {},
                           attempts={a: 1 for a in include})


class TestMermaid:
    def test_module_graph_structure(self, project):
        _, _, _, graph, arch = project
        mermaid = module_graph_mermaid(graph, arch)
        assert mermaid.startswith("flowchart TD")
        assert '"app/main.py"' in mermaid
        assert "-->" in mermaid
        assert "classDef cycle" in mermaid          # cycle styling present
        assert "/" not in mermaid.split("\n")[1].split("[")[0]  # safe ids

    def test_node_cap_with_honest_legend(self):
        edges = [(f"m{i}.py", "core.py") for i in range(60)]
        mermaid = module_graph_mermaid(make_graph(edges), None, max_nodes=10)
        assert mermaid.count("-->") <= 10
        assert "omitted for readability" in mermaid

    def test_waves_subgraphs_and_order(self, project):
        _, _, _, _, arch = project
        mermaid = waves_mermaid(arch)
        assert mermaid.startswith("flowchart LR")
        assert 'subgraph W0["Wave 0' in mermaid
        assert "W0 --> W1" in mermaid


class TestMarkdownReport:
    def test_full_report_with_synthesis(self, project):
        ctx, findings, scorecard, _, _ = project
        synthesis = synthesis_for(findings)
        md = ReportBuilder().build_markdown(ctx, findings, scorecard,
                                            synthesis)

        assert f"# LegacyLens Migration Assessment — proj-demo" in md
        assert f"## Migration readiness: {scorecard.readiness.value}/100" in md
        assert "writer grounded prose" in md
        assert f"[{findings[0].finding_id}]" in md   # citations intact
        assert "```mermaid" in md and "flowchart TD" in md
        assert "## Findings appendix" in md
        for f in findings:
            assert f.finding_id in md                # appendix resolves ids
        assert "heuristic model" in md               # assumptions shipped

    def test_failed_agent_produces_visible_note(self, project):
        ctx, findings, scorecard, _, _ = project
        synthesis = synthesis_for(
            findings, include=("analyst", "writer"),
            failures={"strategist": "citation validation failed after 3 "
                                    "attempts"})
        md = ReportBuilder().build_markdown(ctx, findings, scorecard,
                                            synthesis)
        assert "## Migration plan" in md
        assert "LLM synthesis was unavailable for this section" in md
        assert "citation validation failed" in md
        assert "analyst grounded prose" in md        # others unaffected

    def test_no_synthesis_at_all_still_complete(self, project):
        ctx, findings, scorecard, _, _ = project
        md = ReportBuilder().build_markdown(ctx, findings, scorecard, None)
        assert md.count("LLM synthesis was unavailable") == 3
        assert "## Findings appendix" in md
        assert "## Risk matrix" in md


class TestDocxReport:
    def test_docx_builds_and_reopens(self, project, tmp_path):
        ctx, findings, scorecard, _, _ = project
        out = tmp_path / "report.docx"
        ReportBuilder().build_docx(out, ctx, findings, scorecard,
                                   synthesis_for(findings))

        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs
                    if p.style.name.startswith(("Title", "Heading"))]
        assert any("proj-demo" in h for h in headings)
        assert any("Migration readiness" in h for h in headings)

        all_cells = [c.text for t in doc.tables for r in t.rows
                     for c in r.cells]
        assert findings[0].finding_id in all_cells
        assert "TECH-EOL-001" in all_cells
